"""Export proposal documents to PDF and DOCX (read-only file generation)."""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from typing import Any
from uuid import UUID

from fastapi import HTTPException
from fastapi.responses import Response
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.storage import storage
from app.models.product import Product
from app.models.proposal_document import ProposalDocument
from app.services.proposal_generator_service import SECTION_ORDER

logger = logging.getLogger(__name__)

MARKER = "[Proposal Export]"
FOOTER_TEXT = "Draft commercial proposal"


def _section_labels(language: str) -> dict[str, str]:
    labels = {
        "ru": {
            "intro": "Введение",
            "buyer_need": "Потребность клиента",
            "company_introduction": "О компании",
            "recommended_products": "Рекомендуемые продукты",
            "benefits": "Преимущества",
            "pricing": "Стоимость",
            "moq_payment_delivery": "MOQ, оплата и поставка",
            "next_steps": "Следующие шаги",
            "call_to_action": "Контакты",
        },
        "en": {
            "intro": "Introduction",
            "buyer_need": "Buyer need",
            "company_introduction": "Company introduction",
            "recommended_products": "Recommended products",
            "benefits": "Benefits",
            "pricing": "Pricing",
            "moq_payment_delivery": "MOQ, payment & delivery",
            "next_steps": "Next steps",
            "call_to_action": "Contact",
        },
    }
    return labels.get(language, labels["ru"])


def _strip_markdown(text: str) -> str:
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return text.strip()


@dataclass
class _ExportContext:
    title: str
    language: str
    sections: list[tuple[str, str]]
    client_name: str
    client_block: str
    lead_block: str | None
    deal_block: str | None
    products: list[dict[str, Any]]
    cta: str


class ProposalExportService:
    @staticmethod
    def _storage_folder(proposal_id: UUID) -> str:
        return f"proposals/{proposal_id}"

    @staticmethod
    async def _load_document(db: AsyncSession, proposal_id: UUID) -> ProposalDocument:
        result = await db.execute(
            select(ProposalDocument)
            .options(
                selectinload(ProposalDocument.client),
                selectinload(ProposalDocument.lead),
                selectinload(ProposalDocument.deal),
                selectinload(ProposalDocument.product),
            )
            .where(ProposalDocument.id == proposal_id)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise HTTPException(status_code=404, detail="Proposal document not found")
        return doc

    @staticmethod
    async def _load_products(db: AsyncSession, doc: ProposalDocument) -> list[Product]:
        pj = doc.proposal_json or {}
        raw_ids = pj.get("product_ids") or []
        if not raw_ids and doc.product_id:
            raw_ids = [str(doc.product_id)]
        uuids = []
        for pid in raw_ids:
            try:
                uuids.append(UUID(str(pid)))
            except ValueError:
                continue
        if not uuids:
            return []
        result = await db.execute(select(Product).where(Product.id.in_(uuids)))
        return list(result.scalars().all())

    @staticmethod
    async def _build_context(db: AsyncSession, doc: ProposalDocument) -> _ExportContext:
        pj = doc.proposal_json or {}
        sections_raw = pj.get("sections") or {}
        labels = _section_labels(doc.language)
        sections: list[tuple[str, str]] = []
        for key in SECTION_ORDER:
            body = str(sections_raw.get(key) or "").strip()
            if body:
                sections.append((labels.get(key, key.replace("_", " ").title()), _strip_markdown(body)))

        client = doc.client
        client_lines = [client.company_name]
        if client.business_description:
            client_lines.append(client.business_description[:500])
        if client.cta_phone:
            client_lines.append(f"Phone: {client.cta_phone}")
        if client.cta_telegram:
            client_lines.append(f"Telegram: {client.cta_telegram}")
        if client.cta_website:
            client_lines.append(f"Web: {client.cta_website}")

        lead_block = None
        if doc.lead:
            lead = doc.lead
            parts = [lead.name]
            if lead.company:
                parts.append(lead.company)
            if lead.email:
                parts.append(lead.email)
            if lead.phone:
                parts.append(lead.phone)
            if lead.interest:
                parts.append(f"Interest: {lead.interest[:300]}")
            lead_block = " · ".join(parts)

        deal_block = None
        if doc.deal:
            deal = doc.deal
            deal_block = f"{deal.title} · Status: {deal.status}"
            if deal.expected_value:
                deal_block += f" · Expected: {deal.expected_value}"

        products = await ProposalExportService._load_products(db, doc)
        product_rows = [
            {
                "name": p.name,
                "sku": p.sku or "—",
                "category": p.category or "—",
                "moq": str(p.moq) if p.moq else "TBD",
                "price": f"{p.unit_price} {p.currency}" if p.unit_price else "TBD",
            }
            for p in products
        ]

        cta = sections_raw.get("call_to_action") or ""
        if client.cta_telegram or client.cta_phone:
            cta = cta or f"{client.cta_telegram or client.cta_phone or client.cta_website or ''}"

        return _ExportContext(
            title=doc.title,
            language=doc.language,
            sections=sections,
            client_name=client.company_name,
            client_block="\n".join(client_lines),
            lead_block=lead_block,
            deal_block=deal_block,
            products=product_rows,
            cta=_strip_markdown(str(cta)),
        )

    @staticmethod
    def _build_pdf(ctx: _ExportContext) -> bytes:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2.2 * cm,
            title=ctx.title,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ProposalTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=14,
            textColor=colors.HexColor("#1e293b"),
        )
        h2_style = ParagraphStyle(
            "SectionH",
            parent=styles["Heading2"],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor("#334155"),
        )
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=8,
        )
        meta_style = ParagraphStyle(
            "Meta",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#64748b"),
            spaceAfter=4,
        )

        story: list[Any] = []
        story.append(Paragraph(ctx.title.replace("&", "&amp;"), title_style))
        story.append(Paragraph(f"<b>From:</b> {ctx.client_name.replace('&', '&amp;')}", meta_style))
        if ctx.lead_block:
            story.append(Paragraph(f"<b>Prepared for:</b> {ctx.lead_block.replace('&', '&amp;')}", meta_style))
        if ctx.deal_block:
            story.append(Paragraph(f"<b>Deal:</b> {ctx.deal_block.replace('&', '&amp;')}", meta_style))
        story.append(Spacer(1, 12))

        for heading, body in ctx.sections:
            safe_h = heading.replace("&", "&amp;")
            safe_b = body.replace("&", "&amp;").replace("\n", "<br/>")
            story.append(Paragraph(safe_h, h2_style))
            story.append(Paragraph(safe_b, body_style))

        if ctx.products:
            story.append(Paragraph("Products", h2_style))
            table_data = [["Product", "SKU", "Category", "MOQ", "Price"]]
            for p in ctx.products:
                table_data.append([p["name"], p["sku"], p["category"], p["moq"], p["price"]])
            table = Table(table_data, colWidths=[5 * cm, 2.5 * cm, 3 * cm, 2 * cm, 3 * cm])
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#334155")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ]))
            story.append(table)
            story.append(Spacer(1, 12))

        if ctx.cta:
            story.append(Paragraph("Contact", h2_style))
            story.append(Paragraph(ctx.cta.replace("&", "&amp;").replace("\n", "<br/>"), body_style))

        def _draw_footer(canvas, doc_template):  # noqa: ARG001
            canvas.saveState()
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(colors.HexColor("#94a3b8"))
            canvas.drawCentredString(A4[0] / 2, 1 * cm, FOOTER_TEXT)
            canvas.restoreState()

        doc.build(story, onFirstPage=_draw_footer, onLaterPages=_draw_footer)
        return buffer.getvalue()

    @staticmethod
    def _build_docx(ctx: _ExportContext) -> bytes:
        from docx import Document
        from docx.shared import Pt

        document = Document()
        document.core_properties.title = ctx.title

        document.add_heading(ctx.title, 0)
        p = document.add_paragraph()
        p.add_run("From: ").bold = True
        p.add_run(ctx.client_name)
        if ctx.lead_block:
            p = document.add_paragraph()
            p.add_run("Prepared for: ").bold = True
            p.add_run(ctx.lead_block)
        if ctx.deal_block:
            p = document.add_paragraph()
            p.add_run("Deal: ").bold = True
            p.add_run(ctx.deal_block)

        for heading, body in ctx.sections:
            document.add_heading(heading, level=1)
            for para in body.split("\n"):
                if para.strip():
                    document.add_paragraph(para.strip())

        if ctx.products:
            document.add_heading("Products", level=1)
            table = document.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, label in enumerate(["Product", "SKU", "Category", "MOQ", "Price"]):
                hdr[i].text = label
                for run in hdr[i].paragraphs[0].runs:
                    run.bold = True
            for prod in ctx.products:
                row = table.add_row().cells
                row[0].text = prod["name"]
                row[1].text = prod["sku"]
                row[2].text = prod["category"]
                row[3].text = prod["moq"]
                row[4].text = prod["price"]

        if ctx.cta:
            document.add_heading("Contact", level=1)
            document.add_paragraph(ctx.cta)

        footer = document.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()
        run = footer_para.add_run(FOOTER_TEXT)
        run.font.size = Pt(8)

        bio = BytesIO()
        document.save(bio)
        return bio.getvalue()

    @staticmethod
    async def export_pdf(db: AsyncSession, proposal_id: UUID) -> dict[str, Any]:
        logger.info("%s pdf started: id=%s", MARKER, proposal_id)
        doc = await ProposalExportService._load_document(db, proposal_id)
        try:
            ctx = await ProposalExportService._build_context(db, doc)
            pdf_bytes = ProposalExportService._build_pdf(ctx)
            folder = ProposalExportService._storage_folder(proposal_id)
            key = f"{folder}/proposal.pdf"
            await storage.save_at_key(key, pdf_bytes)
            now = datetime.now(timezone.utc)
            doc.exported_pdf_path = key
            doc.last_exported_at = now
            doc.updated_at = now
            await db.commit()
            await db.refresh(doc)
            logger.info("%s completed: id=%s format=pdf path=%s", MARKER, proposal_id, key)
            return {
                "id": doc.id,
                "format": "pdf",
                "path": key,
                "last_exported_at": now,
                "download_url": storage.get_url(key),
            }
        except HTTPException:
            raise
        except Exception as exc:
            logger.exception("%s failed: id=%s format=pdf error=%s", MARKER, proposal_id, exc)
            raise HTTPException(
                status_code=500,
                detail=f"PDF export failed: {exc}",
            ) from exc

    @staticmethod
    async def export_docx(db: AsyncSession, proposal_id: UUID) -> dict[str, Any]:
        logger.info("%s docx started: id=%s", MARKER, proposal_id)
        doc = await ProposalExportService._load_document(db, proposal_id)
        try:
            ctx = await ProposalExportService._build_context(db, doc)
            docx_bytes = ProposalExportService._build_docx(ctx)
            folder = ProposalExportService._storage_folder(proposal_id)
            key = f"{folder}/proposal.docx"
            await storage.save_at_key(key, docx_bytes)
            now = datetime.now(timezone.utc)
            doc.exported_docx_path = key
            doc.last_exported_at = now
            doc.updated_at = now
            await db.commit()
            await db.refresh(doc)
            logger.info("%s completed: id=%s format=docx path=%s", MARKER, proposal_id, key)
            return {
                "id": doc.id,
                "format": "docx",
                "path": key,
                "last_exported_at": now,
                "download_url": storage.get_url(key),
            }
        except HTTPException:
            raise
        except Exception as exc:
            logger.exception("%s failed: id=%s format=docx error=%s", MARKER, proposal_id, exc)
            raise HTTPException(
                status_code=500,
                detail=f"DOCX export failed: {exc}",
            ) from exc

    @staticmethod
    async def download_pdf(db: AsyncSession, proposal_id: UUID) -> Response:
        doc = await ProposalExportService._load_document(db, proposal_id)
        if not doc.exported_pdf_path or not storage.exists(doc.exported_pdf_path):
            raise HTTPException(status_code=404, detail="PDF export not found — run export first")
        try:
            data = await storage.read_file_bytes(doc.exported_pdf_path)
        except FileNotFoundError:
            raise HTTPException(status_code=404, detail="PDF file missing on storage") from None
        filename = f"proposal-{proposal_id}.pdf"
        return Response(
            content=data,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    @staticmethod
    async def download_docx(db: AsyncSession, proposal_id: UUID) -> Response:
        doc = await ProposalExportService._load_document(db, proposal_id)
        if not doc.exported_docx_path or not storage.exists(doc.exported_docx_path):
            raise HTTPException(status_code=404, detail="DOCX export not found — run export first")
        try:
            data = await storage.read_file_bytes(doc.exported_docx_path)
        except FileNotFoundError:
            raise HTTPException(status_code=404, detail="DOCX file missing on storage") from None
        filename = f"proposal-{proposal_id}.docx"
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
